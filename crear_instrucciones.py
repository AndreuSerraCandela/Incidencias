"""
Script para crear un documento Word con instrucciones de uso de la aplicación
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Crear documento
doc = Document()

# Configurar título principal
title = doc.add_heading('Instrucciones de Uso - Aplicación de Incidencias', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Agregar introducción
intro = doc.add_paragraph()
intro.add_run('Esta aplicación permite reportar incidencias en paradas de autobús utilizando tres métodos diferentes:').bold = True
doc.add_paragraph('1. Escaneo de etiquetas NFC')
doc.add_paragraph('2. Grabación de audio')
doc.add_paragraph('3. Procesamiento de imágenes con IA')

doc.add_paragraph()  # Espacio

# ============================================
# SECCIÓN 1: NFC (Etiquetas NFC)
# ============================================
doc.add_heading('1. Reportar Incidencia con NFC', 1)

doc.add_paragraph('El escaneo NFC es automático y continuo. La aplicación escanea etiquetas NFC de forma automática cuando estás logueado.')

doc.add_heading('Pasos a seguir:', 2)
doc.add_paragraph('1. Inicia sesión en la aplicación con tus credenciales.', style='List Number')
doc.add_paragraph('2. Espera a que aparezcan los botones de acción. El escaneo NFC se activará automáticamente.', style='List Number')
doc.add_paragraph('3. Acerca tu dispositivo a la etiqueta NFC de la parada de autobús.', style='List Number')
doc.add_paragraph('4. Cuando la etiqueta NFC se lea correctamente:', style='List Number')

sub_list = doc.add_paragraph()
sub_list.add_run('   • Se mostrará un mensaje de éxito').italic = True
sub_list = doc.add_paragraph()
sub_list.add_run('   • Se mostrarán los datos de la etiqueta NFC en pantalla').italic = True
sub_list = doc.add_paragraph()
sub_list.add_run('   • Se abrirá automáticamente la cámara para capturar una foto').italic = True

doc.add_paragraph('5. Captura una foto de la incidencia usando el botón "Capturar Foto" o "Importar Foto".', style='List Number')
doc.add_paragraph('6. Revisa la foto y haz clic en "Enviar Incidencia" para completar el reporte.', style='List Number')

doc.add_heading('Notas importantes:', 2)
note1 = doc.add_paragraph()
note1.add_run('• El escaneo NFC funciona automáticamente: ').bold = True
note1.add_run('No necesitas hacer clic en ningún botón. Solo acerca el dispositivo a la etiqueta.')
doc.add_paragraph('• El escaneo se detiene temporalmente cuando pulsas "Grabar Audio" o "Reportar Incidencia".')
doc.add_paragraph('• El escaneo se reinicia automáticamente después de enviar una incidencia.')
doc.add_paragraph('• Algunos dispositivos o navegadores pueden no soportar NFC. En ese caso, verás un mensaje de advertencia.')

doc.add_page_break()

# ============================================
# SECCIÓN 2: Audio
# ============================================
doc.add_heading('2. Reportar Incidencia con Audio', 1)

doc.add_paragraph('Puedes grabar un mensaje de voz describiendo la incidencia y la aplicación lo transcribirá automáticamente.')

doc.add_heading('Pasos a seguir:', 2)
doc.add_paragraph('1. Haz clic en el botón "Grabar Audio".', style='List Number')
doc.add_paragraph('2. En la ventana que aparece, haz clic en "Iniciar Grabación".', style='List Number')
doc.add_paragraph('3. Permite el acceso al micrófono cuando el navegador lo solicite.', style='List Number')
doc.add_paragraph('4. Habla claramente describiendo:', style='List Number')

sub_list = doc.add_paragraph()
sub_list.add_run('   • El número de parada (ejemplo: "Parada 625" o "Parada P625")').italic = True
sub_list = doc.add_paragraph()
sub_list.add_run('   • La descripción de la incidencia (ejemplo: "Cristal roto", "Pintada", "Grafiti")').italic = True

doc.add_paragraph('5. Cuando termines, haz clic en "Detener Grabación".', style='List Number')
doc.add_paragraph('6. Puedes reproducir el audio para verificar que se grabó correctamente.', style='List Number')
doc.add_paragraph('7. Si estás satisfecho, haz clic en "Usar Audio".', style='List Number')
doc.add_paragraph('8. La aplicación procesará el audio y extraerá automáticamente:', style='List Number')

sub_list = doc.add_paragraph()
sub_list.add_run('   • El número de parada').italic = True
sub_list = doc.add_paragraph()
sub_list.add_run('   • La descripción de la incidencia').italic = True

doc.add_paragraph('9. Revisa los resultados y, si es necesario, captura una foto adicional usando "Reportar Incidencia".', style='List Number')
doc.add_paragraph('10. Haz clic en "Enviar Incidencia" para completar el reporte.', style='List Number')

doc.add_heading('Consejos para una mejor transcripción:', 2)
doc.add_paragraph('• Habla claramente y a un ritmo moderado.')
doc.add_paragraph('• Reduce el ruido de fondo lo máximo posible.')
doc.add_paragraph('• Menciona el número de parada al inicio del mensaje (ejemplo: "Parada 625, el cristal está roto").')
doc.add_paragraph('• Usa números simples en lugar de palabras complejas (di "625" en lugar de "seiscientos veinticinco").')

doc.add_heading('Notas importantes:', 2)
doc.add_paragraph('• El procesamiento de audio puede tardar unos segundos.')
doc.add_paragraph('• La aplicación usa inteligencia artificial para extraer la información del audio.')
doc.add_paragraph('• Puedes grabar nuevamente si no estás satisfecho con el resultado.')
doc.add_paragraph('• El audio se procesa en el servidor y no se almacena permanentemente.')

doc.add_page_break()

# ============================================
# SECCIÓN 3: IA Imagen
# ============================================
doc.add_heading('3. Reportar Incidencia con Procesamiento de Imagen con IA', 1)

doc.add_paragraph('Puedes capturar o subir una foto de la parada y la aplicación usará inteligencia artificial para identificar automáticamente el número de parada y la incidencia.')

doc.add_heading('Pasos a seguir:', 2)
doc.add_paragraph('1. Haz clic en el botón "Reportar Incidencia".', style='List Number')
doc.add_paragraph('2. En la ventana que aparece, tienes dos opciones:', style='List Number')

sub_list = doc.add_paragraph()
sub_list.add_run('   a) Capturar Foto: Usa la cámara de tu dispositivo para tomar una foto').italic = True
sub_list = doc.add_paragraph()
sub_list.add_run('   b) Importar Foto: Selecciona una foto existente de tu dispositivo').italic = True

doc.add_paragraph('3. Asegúrate de que la foto muestre claramente:', style='List Number')

sub_list = doc.add_paragraph()
sub_list.add_run('   • El cartel informativo de la parada con el número de parada (código que empieza con "P")').italic = True
sub_list = doc.add_paragraph()
sub_list.add_run('   • La incidencia o pintada visible (si existe)').italic = True

doc.add_paragraph('4. Una vez capturada o importada la foto, haz clic en "Subir Foto".', style='List Number')
doc.add_paragraph('5. La aplicación procesará la imagen con inteligencia artificial. Esto puede tardar unos momentos.', style='List Number')
doc.add_paragraph('6. Se mostrará un modal con los resultados de la IA:', style='List Number')

sub_list = doc.add_paragraph()
sub_list.add_run('   • Número de parada detectado').italic = True
sub_list = doc.add_paragraph()
sub_list.add_run('   • Descripción de la incidencia detectada').italic = True

doc.add_paragraph('7. Revisa los resultados. Puedes modificar el número de parada o la descripción si es necesario.', style='List Number')
doc.add_paragraph('8. Cuando estés satisfecho, haz clic en "Confirmar Resultados".', style='List Number')
doc.add_paragraph('9. Haz clic en "Enviar Incidencia" para completar el reporte.', style='List Number')

doc.add_heading('Consejos para una mejor detección:', 2)
doc.add_paragraph('• Asegúrate de que la foto tenga buena iluminación.')
doc.add_paragraph('• El número de parada debe ser visible y legible en la foto.')
doc.add_paragraph('• Evita fotos borrosas o movidas.')
doc.add_paragraph('• Si la incidencia es visible, captúrala en la misma foto.')
doc.add_paragraph('• No confundas el número de parada (código con "P") con el número de línea del autobús (que aparece en carteles de color).')

doc.add_heading('Notas importantes:', 2)
doc.add_paragraph('• El procesamiento de imagen con IA puede tardar entre 10 y 30 segundos.')
doc.add_paragraph('• La IA analiza la imagen y busca automáticamente el número de parada y las incidencias visibles.')
doc.add_paragraph('• Puedes corregir manualmente los resultados si la IA no los detecta correctamente.')
doc.add_paragraph('• Si la foto no muestra el número de parada, puedes introducirlo manualmente.')

doc.add_page_break()

# ============================================
# SECCIÓN 4: Información General
# ============================================
doc.add_heading('Información General', 1)

doc.add_heading('Requisitos del dispositivo:', 2)
doc.add_paragraph('• Navegador web moderno (Chrome, Firefox, Edge, Safari)')
doc.add_paragraph('• Conexión a Internet')
doc.add_paragraph('• Para NFC: Dispositivo con soporte NFC y navegador compatible')
doc.add_paragraph('• Para Audio: Micrófono funcional')
doc.add_paragraph('• Para Imágenes: Cámara o capacidad de subir archivos de imagen')

doc.add_heading('Permisos necesarios:', 2)
doc.add_paragraph('• Cámara: Para capturar fotos de las incidencias')
doc.add_paragraph('• Micrófono: Para grabar audio')
doc.add_paragraph('• NFC: Para leer etiquetas NFC (si está disponible)')

doc.add_heading('Solución de problemas:', 2)

doc.add_heading('No aparece el botón de NFC:', 3)
doc.add_paragraph('El escaneo NFC es automático y no requiere botón. Si no funciona, verifica que tu dispositivo y navegador soporten NFC.')

doc.add_heading('El audio no se graba:', 3)
doc.add_paragraph('Verifica que hayas permitido el acceso al micrófono en la configuración del navegador.')

doc.add_heading('La cámara no funciona:', 3)
doc.add_paragraph('Verifica que hayas permitido el acceso a la cámara y que no esté siendo usada por otra aplicación.')

doc.add_heading('La IA no detecta el número de parada:', 3)
doc.add_paragraph('Asegúrate de que la foto muestre claramente el cartel informativo con el código de parada. Puedes introducir el número manualmente si es necesario.')

# Guardar documento
doc.save('Instrucciones_Uso_Aplicacion_Incidencias.docx')
print('Documento Word creado exitosamente: Instrucciones_Uso_Aplicacion_Incidencias.docx')

