"""
Genera Manual de Uso de Incidencias (Word) con pantallazos,
estilo similar a Manuales/Tareas.docx.
"""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).resolve().parent
SHOTS = BASE / "manual_screenshots"
OUT_PROJECT = BASE.parent / "Instrucciones_Uso_Aplicacion_Incidencias.docx"
OUT_MANUALES = Path(r"c:\Users\Andres\OneDrive\Documentos\Manuales\Incidencias.docx")


def add_shot(doc, filename, width_in=5.2, caption=None):
    path = SHOTS / filename
    if not path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Falta captura: {filename}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_note(doc, text):
    p = doc.add_paragraph()
    p.add_run("Nota: ").bold = True
    p.add_run(text)


def build():
    doc = Document()

    title = doc.add_heading("Incidencias", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("(Aplicación web / PWA — Grupo Malla)")
    r.italic = True

    doc.add_heading("Introducción", level=1)
    p = doc.add_paragraph()
    r = p.add_run("(objetivo del documento)")
    r.italic = True
    doc.add_paragraph(
        "Este documento orienta al usuario en el uso de la aplicación de Incidencias: "
        "inicio de sesión, reporte de incidencias (foto, audio, NFC, mapa), "
        "cierre de incidencias abiertas, y el flujo de guardar fotos en el servidor "
        "para asignarlas después (campo u oficina)."
    )

    doc.add_heading("Pantalla principal", level=1)
    doc.add_paragraph(
        "Tras iniciar sesión verás los botones de acción. En móvil la disposición es "
        "en dos columnas; en escritorio se ve de forma similar, más ancha."
    )
    add_shot(doc, "01_inicio_login.png", caption="Figura 1. Acceso / inicio de sesión")
    add_shot(doc, "02_modal_login.png", caption="Figura 2. Modal de credenciales GTask")
    add_shot(doc, "03_botones_principales.png", caption="Figura 3. Botones principales (móvil)")
    add_shot(doc, "12_escritorio_botones.png", width_in=5.8, caption="Figura 4. Vista en escritorio")

    doc.add_paragraph("Botones disponibles:")
    doc.add_paragraph("Grabar Audio — graba un mensaje de voz y lo transcribe.", style="List Bullet")
    doc.add_paragraph("Elementos Cerca — mapa con recursos/mobiliario cercanos.", style="List Bullet")
    doc.add_paragraph("Añadir Fotos — fotos adicionales a una incidencia en curso.", style="List Bullet")
    doc.add_paragraph("Reportar Incidencia — captura foto y envía incidencia nueva.", style="List Bullet")
    doc.add_paragraph("Tomar Foto — guarda foto en el servidor (con GPS y descripción opcional).", style="List Bullet")
    doc.add_paragraph("Asignar foto — elige una foto guardada y la asocia a un elemento del mapa.", style="List Bullet")

    doc.add_page_break()

    # ---- Reportar ----
    doc.add_heading("Reportar una incidencia", level=1)
    doc.add_heading("Con foto (Reportar Incidencia)", level=2)
    doc.add_paragraph("1. Pulsa Reportar Incidencia.", style="List Number")
    doc.add_paragraph("2. Captura o importa una foto.", style="List Number")
    doc.add_paragraph("3. Revisa la vista previa y pulsa Enviar Incidencia.", style="List Number")
    doc.add_paragraph(
        "4. Elige tipo/subtipo y escribe la descripción (puede venir pre-rellenada si hubo audio, IA o descripción de foto local).",
        style="List Number",
    )
    add_shot(doc, "04_modal_camara.png", caption="Figura 5. Modal de captura / importar foto")
    add_shot(doc, "05_vista_previa_enviar.png", caption="Figura 6. Vista previa y botón Enviar")
    add_shot(doc, "11_tipo_subtipo.png", caption="Figura 7. Selección de tipo y subtipo")

    doc.add_heading("Con audio", level=2)
    doc.add_paragraph(
        "Pulsa Grabar Audio, inicia la grabación, describe la parada y la incidencia, "
        "detén y pulsa Usar Audio. Después completa con foto si hace falta."
    )
    add_shot(doc, "06_modal_audio.png", caption="Figura 8. Grabación de audio")
    add_note(
        doc,
        "La transcripción por Whisper requiere que esté instalado en el servidor "
        "(install_whisper.bat y FFmpeg). Si no está disponible, el servidor devolverá error.",
    )

    doc.add_heading("Con Elementos Cerca", level=2)
    doc.add_paragraph("1. Pulsa Elementos Cerca y acepta el permiso de ubicación.", style="List Number")
    doc.add_paragraph(
        "2. Se muestra un mapa con tu posición (o la de la foto), el radio de búsqueda y los elementos.",
        style="List Number",
    )
    doc.add_paragraph(
        "3. Puedes ampliar el radio con el botón del encabezado, o arrastrar el pin / tocar el mapa si la ubicación del PC es imprecisa.",
        style="List Number",
    )
    doc.add_paragraph(
        "4. En el popup del elemento: Crear Incidencia (abre cámara) o Cerrar incidencia (si hay una abierta tuya).",
        style="List Number",
    )
    add_shot(doc, "07_elementos_cerca.png", caption="Figura 9. Modal Elementos Cerca")
    add_note(
        doc,
        "En PC el GPS del navegador a menudo es poco preciso. Arrastra el marcador azul "
        "o usa «Última ubicación buena» si ya guardaste una posición fiable (p. ej. desde el móvil).",
    )

    doc.add_heading("NFC", level=2)
    doc.add_paragraph(
        "Con sesión iniciada, el escaneo NFC puede activarse automáticamente en dispositivos compatibles. "
        "Al leer una etiqueta se prepara la incidencia y se abre la cámara."
    )

    doc.add_page_break()

    # ---- Tomar / Asignar ----
    doc.add_heading("Tomar Foto y Asignar foto", level=1)
    p = doc.add_paragraph()
    r = p.add_run("(flujo campo → oficina)")
    r.italic = True
    doc.add_paragraph(
        "Sirve para capturar en campo y asignar el recurso después (en el mismo dispositivo o desde otro PC/oficina). "
        "Las fotos se guardan en el servidor asociadas al usuario GTask."
    )

    doc.add_heading("Tomar Foto", level=2)
    doc.add_paragraph("1. Pulsa Tomar Foto y captura o importa la imagen.", style="List Number")
    doc.add_paragraph(
        "2. (Opcional) Añade una descripción escrita o con Dictar (reconocimiento de voz del navegador).",
        style="List Number",
    )
    doc.add_paragraph("3. Guarda u omite la descripción: la foto queda en el servidor con tu usuario y, si hay GPS, la ubicación.", style="List Number")
    add_shot(doc, "08_descripcion_foto_local.png", caption="Figura 10. Descripción opcional al guardar la foto")

    doc.add_heading("Asignar foto", level=2)
    doc.add_paragraph("1. Pulsa Asignar foto.", style="List Number")
    doc.add_paragraph("2. Por defecto ves Solo mías.", style="List Number")
    doc.add_paragraph(
        "3. Pulsa Todos los usuarios para ver las fotos de todos; cada usuario se puede plegar/desplegar.",
        style="List Number",
    )
    doc.add_paragraph("4. Elige una foto → se abre Elementos cerca en la ubicación de la foto.", style="List Number")
    doc.add_paragraph("5. En un recurso, Crear Incidencia: se carga esa foto (sin volver a abrir la cámara).", style="List Number")
    doc.add_paragraph("6. Completa tipo/descripción (pre-rellenada si la foto tenía descripción) y envía.", style="List Number")
    doc.add_paragraph(
        "7. Si el envío es correcto, la foto se elimina del servidor.",
        style="List Number",
    )
    add_shot(doc, "09_asignar_fotos_mias.png", caption="Figura 11. Listado «Solo mías»")
    add_shot(doc, "10_asignar_fotos_todos.png", caption="Figura 12. Listado «Todos los usuarios» (acordeón)")

    doc.add_page_break()

    # ---- Cerrar ----
    doc.add_heading("Cerrar una incidencia abierta", level=1)
    doc.add_paragraph(
        "Desde Elementos Cerca, si un elemento tiene incidencia abierta asignada a ti, "
        "aparece Cerrar incidencia. Se abre la cámara para la foto de cierre; al pulsar "
        "Cerrar incidencia se envía el estado Cerrada (y EnProgreso si aplica) con el número de documento BC."
    )
    doc.add_paragraph(
        "También puedes abrir una incidencia por enlace (p. ej. ?id=INC000750): se muestra un resumen "
        "y puedes continuar a la foto de cierre."
    )

    doc.add_heading("Consejos y limitaciones", level=1)
    doc.add_paragraph("Usa el móvil con GPS para Tomar Foto si luego vas a asignar en oficina.", style="List Bullet")
    doc.add_paragraph("En PC, corrige el centro del mapa arrastrando el pin cuando la precisión sea mala.", style="List Bullet")
    doc.add_paragraph("Las fotos pendientes viven en el servidor (carpeta pending_photos), no en el navegador.", style="List Bullet")
    doc.add_paragraph("Tras enviar la incidencia con una foto asignada, esa foto se borra automáticamente.", style="List Bullet")
    doc.add_paragraph("Recarga forzada (Ctrl+F5) si no ves cambios recientes de la aplicación.", style="List Bullet")

    doc.add_heading("Resumen de pantallas", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Pantalla"
    hdr[1].text = "Uso"
    rows = [
        ("Inicio / Login", "Acceso con usuario GTask"),
        ("Botones principales", "Elegir el flujo de trabajo"),
        ("Cámara", "Capturar o importar foto"),
        ("Vista previa", "Revisar y enviar"),
        ("Audio", "Dictar incidencia"),
        ("Elementos Cerca", "Mapa, ampliar radio, crear/cerrar"),
        ("Descripción foto", "Texto/dictado opcional al guardar"),
        ("Asignar foto", "Mis fotos o todas; asignar a recurso"),
        ("Tipo / subtipo", "Clasificar la incidencia"),
    ]
    for a, b in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b

    # Guardar
    doc.save(str(OUT_PROJECT))
    print(f"Guardado: {OUT_PROJECT}")
    try:
        OUT_MANUALES.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(OUT_MANUALES))
        print(f"Guardado: {OUT_MANUALES}")
    except Exception as e:
        print(f"No se pudo guardar en Manuales: {e}")


if __name__ == "__main__":
    build()
